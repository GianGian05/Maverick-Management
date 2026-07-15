from __future__ import annotations

import hashlib
import json
import shutil
import sys
from pathlib import Path

from docx import Document


def replace_in_paragraph(paragraph, old: str, new: str) -> int:
    """Replace text across Word runs while preserving surrounding formatting."""
    count = 0
    while True:
        runs = paragraph.runs
        full = "".join(run.text for run in runs)
        start = full.find(old)
        if start < 0:
            return count
        end = start + len(old)
        positions = []
        cursor = 0
        for idx, run in enumerate(runs):
            positions.append((idx, cursor, cursor + len(run.text)))
            cursor += len(run.text)
        start_run = next((p for p in positions if p[1] <= start < p[2]), None)
        end_run = next((p for p in positions if p[1] < end <= p[2]), None)
        if start_run is None or end_run is None:
            return count
        si, ss, _ = start_run
        ei, es, _ = end_run
        prefix = runs[si].text[: start - ss]
        suffix = runs[ei].text[end - es :]
        runs[si].text = prefix + new + suffix
        for idx in range(si + 1, ei + 1):
            runs[idx].text = ""
        count += 1


def all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: prepare_rfp_template.py SOURCE.docx OUTPUT_DIR")
    source = Path(sys.argv[1]).resolve()
    output_dir = Path(sys.argv[2]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / "Roof_RFP_Master_Template.docx"
    shutil.copy2(source, output)

    replacements = {
        "Copy to: [lspongberg@trustmaverick.com], [tgrimes@trustmaverick.com], [yliu@trustmaverick.com]": "Copy to: {{COPY_EMAILS}}",
        "Base Bid – Option B - 20-year 60mil TPO Mechanically Fastened": "Base Bid – Option B - {{OPTION_B_NAME}}",
        "Base Bid – Option A - 20-year 60mil TPO RhinoBond": "Base Bid – Option A - {{OPTION_A_NAME}}",
        "6820 Shingle Creek Pkwy, Minneapolis, MN 55430": "{{FULL_ADDRESS}}",
        "6820 Shingle Creek Pkwy 75212": "{{ADDRESS_LINE_1}}",
        "Minneapolis, MN 55430": "{{CITY_STATE_ZIP}}",
        "GKI Industrial Minneapolis, LLC": "{{OWNER_ENTITY_LEGAL}}",
        "GKI Industrial Minneapolis": "{{OWNER_ENTITY}}",
        "[March 23, 2026]": "{{RFP_ISSUE_DATE_LONG}}",
        "March 23, 2026": "{{RFP_ISSUE_DATE_LONG}}",
        "[3/23/2026]": "{{RFP_ISSUE_DATE_SHORT}}",
        "[Jan 30, 2026, 3PM]": "{{BID_DUE}}",
        "[Schedule with MMG]": "{{JOB_WALK}}",
        "[146,000]": "{{AREA_SF}}",
        "State of [Georgia]": "State of {{LICENSE_STATE}}",
        "State of [Minnesota]": "State of {{LICENSE_STATE}}",
        "[Georgia]": "{{LICENSE_STATE}}",
        "[Minnesota]": "{{LICENSE_STATE}}",
        "[fredtuck@trustmaverick.com]": "{{PRIMARY_CONTACT_EMAIL}}",
        "fredtuck@trustmaverick.com": "{{PRIMARY_CONTACT_EMAIL}}",
        "Fred Tuck": "{{PRIMARY_CONTACT_NAME}}",
        "$5,000,000.00": "{{LIABILITY_LIMIT}}",
        "ROOF Recover": "{{ROOF_PROJECT_TYPE}}",
        "20-year 60mil TPO RhinoBond": "{{OPTION_A_NAME}}",
        "20-year 60mil TPO Mechanically Fastened": "{{OPTION_B_NAME}}",
    }

    doc = Document(output)
    counts = {}
    for old, new in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        total = 0
        for paragraph in all_paragraphs(doc):
            total += replace_in_paragraph(paragraph, old, new)
        counts[old] = total
    doc.save(output)

    manifest = {
        "template_version": "1.0",
        "template_file": output.name,
        "sha256": sha256(output),
        "replacement_counts": counts,
        "source_file": source.name,
        "locked_language_policy": "Only placeholders may be populated. All other template content is treated as locked.",
    }
    (output_dir / "RFP_Template_Manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
